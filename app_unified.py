import gradio as gr
import whisper
import ollama
import cv2
import os
import base64
import shutil
import atexit
from datetime import timedelta
from xhtml2pdf import pisa
from docx import Document

# --- CONFIGURATION ---
MODEL_NAME = "gemma3:4b"  
TEMP_DIR = "temp_unified" 

# Mapping for Whisper Language Codes
LANG_MAP = {
    "Auto-Detect": None,
    "English": "en",
    "Spanish": "es",
    "Japanese": "ja",
    "French": "fr",
    "German": "de",
    "Italian": "it",
    "Portuguese": "pt",
    "Chinese": "zh"
}

os.makedirs(TEMP_DIR, exist_ok=True)

# --- CLEANUP ---
def cleanup_on_exit():
    if os.path.exists(TEMP_DIR):
        try: shutil.rmtree(TEMP_DIR)
        except: pass

atexit.register(cleanup_on_exit)

# --- HELPER FUNCTIONS ---
def image_to_base64(image):
    _, buffer = cv2.imencode('.jpg', image)
    return base64.b64encode(buffer).decode('utf-8')

def extract_frame_base64(video_path, seconds):
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened(): return None
    cap.set(cv2.CAP_PROP_POS_MSEC, seconds * 1000)
    success, image = cap.read()
    cap.release()
    if success:
        image = cv2.resize(image, (640, 360)) 
        return image_to_base64(image)
    return None

def is_video_file(filename):
    video_exts = ('.mp4', '.mkv', '.mov', '.avi', '.webm')
    return filename.lower().endswith(video_exts)

def convert_html_to_pdf(source_html, output_filename):
    try:
        with open(output_filename, "wb") as result_file:
            pisa_status = pisa.CreatePDF(source_html, dest=result_file)
        return not pisa_status.err
    except Exception as e:
        print(f"PDF Gen Error: {e}")
        return False

def chunk_text(text, chunk_size=6000, overlap=100):
    """Splits text into chunks of roughly chunk_size with overlap."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = start + chunk_size
        
        # adjusting end to avoid cutting words if possible (optional simple logic)
        if end < text_len:
            # find last space within the lookahead buffer handling
            last_space = text.rfind(' ', start, end)
            if last_space != -1 and last_space > start + (chunk_size // 2):
                end = last_space
        
        chunks.append(text[start:end])
        start = end - overlap # move back for overlap
        
        # Prevent infinite loop if overlap >= chunk_size or similar edge cases
        if start >= end:
            start = end
            
    return chunks

def save_as_docx(title, summary, output_filename):
    try:
        doc = Document()
        doc.add_heading(title, 0)
        
        doc.add_heading('AI Output', level=1)
        # Parse <br> for newlines
        lines = summary.split("<br>")
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())
        
        doc.save(output_filename)
        return True
    except Exception as e:
        print(f"DOCX Error: {e}")
        return False

# ==========================================
# LOGIC 1: MEETING ASSISTANT (Standard)
# ==========================================
def process_assistant(file_obj, keyword, custom_prompt, progress=gr.Progress()):
    if file_obj is None: return "Please upload a file.", None, None
    
    file_path = file_obj.name
    filename = os.path.basename(file_path)
    file_prefix = os.path.splitext(filename)[0]

    # 1. Transcribe (Auto language)
    progress(0.1, desc="Transcribing...")
    model = whisper.load_model("base")
    result = model.transcribe(file_path, fp16=False)
    full_text = result['text']
    segments = result['segments']

    # 2. Mentions
    progress(0.4, desc="Searching Keyword...")
    mentions_html = ""
    found_count = 0
    for segment in segments:
        if keyword.lower() in segment['text'].lower():
            found_count += 1
            time_str = str(timedelta(seconds=int(segment['start'])))
            img_tag = ""
            if is_video_file(file_path):
                b64 = extract_frame_base64(file_path, segment['start'])
                if b64: img_tag = f'<br><img src="data:image/jpeg;base64,{b64}" style="width:400px;">'
            mentions_html += f"<div class='mention'><strong>{time_str}</strong>: {segment['text']}{img_tag}</div>"

    # 3. Summarize
    progress(0.7, desc="AI Summarizing (Chunking)...")
    hidden_rules = "STRICT: Output HTML-friendly text. No fillers. No questions."
    
    text_chunks = chunk_text(full_text)
    summary_parts = []

    for i, chunk in enumerate(text_chunks):
        progress((0.7 + (0.2 * (i+1)/len(text_chunks))), desc=f"Summarizing chunk {i+1}/{len(text_chunks)}...")
        final_prompt = f"{hidden_rules}\n{custom_prompt}\nTRANSCRIPT PART {i+1}:\n{chunk}"
        
        try:
            response = ollama.chat(model=MODEL_NAME, messages=[{'role': 'user', 'content': final_prompt}])
            part_content = response['message']['content'].replace("\n", "<br>")
            summary_parts.append(part_content)
        except Exception as e:
            summary_parts.append(f"<i>Error in chunk {i+1}: {e}</i>")

    summary = "<br><hr><br>".join(summary_parts)

    # 4. Render
    html_out = f"<h1>Report: {filename}</h1><h2>Summary</h2><div class='box'>{summary}</div><h2>Mentions ({found_count})</h2>{mentions_html}"
    
    # Save Files
    h_path = os.path.join(TEMP_DIR, f"Assist_{file_prefix}.html")
    p_path = os.path.join(TEMP_DIR, f"Assist_{file_prefix}.pdf")
    
    # Full HTML wrapper for PDF
    full_html = f"<html><body>{html_out}</body></html>"
    with open(h_path, "w", encoding="utf-8") as f: f.write(full_html)
    convert_html_to_pdf(full_html, p_path)

    return full_html, p_path, [h_path, p_path]


# ==========================================
# LOGIC 2: UNIVERSAL TRANSLATOR (Dynamic)
# ==========================================
def process_translator(file_obj, source_lang, target_lang, progress=gr.Progress()):
    if file_obj is None: return "Please upload a file.", None

    file_path = file_obj.name
    filename = os.path.basename(file_path)
    file_prefix = os.path.splitext(filename)[0]
    
    # 1. Transcribe with Source Language
    progress(0.1, desc=f"Transcribing ({source_lang})...")
    model = whisper.load_model("base")
    
    whisper_lang_code = LANG_MAP.get(source_lang) # Convert "Spanish" to "es"
    
    # If None, Whisper auto-detects
    result = model.transcribe(file_path, fp16=False, language=whisper_lang_code)
    
    full_text_original = result['text']
    segments = result['segments']

    # 2. Mentions (SKIPPED based on user request)
    # User requested only AI Output and PDF save.


    # 3. Translate & Summarize
    # 3. Translate & Summarize
    progress(0.7, desc=f"Translating to {target_lang} (Chunking)...")
    
    translation_prompt = f"""
    ROLE: Professional Interpreter.
    TASK: Translate the following transcript from {source_lang} into {target_lang}.
    INSTRUCTIONS:
    1. Provide a clear and accurate translation of the segment.
    2. Output strictly in {target_lang}.
    3. Do not add conversational fillers.
    4. Maintain continuity if possible.
    """
    
    text_chunks = chunk_text(full_text_original)
    translated_parts = []
    
    for i, chunk in enumerate(text_chunks):
        progress((0.7 + (0.2 * (i+1)/len(text_chunks))), desc=f"Translating chunk {i+1}/{len(text_chunks)}...")
        
        final_prompt = f"{translation_prompt}\n\nORIGINAL TRANSCRIPT PART {i+1}:\n{chunk}"
        
        try:
            response = ollama.chat(model=MODEL_NAME, messages=[{'role': 'user', 'content': final_prompt}])
            part_content = response['message']['content'].replace("\n", "<br>")
            translated_parts.append(part_content)
        except Exception as e:
            translated_parts.append(f"<i>Error in chunk {i+1}: {e}</i>")

    summary = "<br><hr><br>".join(translated_parts)

    # 4. Render HTML (No PDF for translation to avoid font issues)
    # 4. Render HTML
    html_out = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: sans-serif; background: #0b0f19; color: #e5e7eb; padding: 20px; }}
            h1 {{ color: #60a5fa; border-bottom: 1px solid #333; }}
            .box {{ background: #1f2937; padding: 20px; border-radius: 8px; margin-top: 10px; }}
        </style>
    </head>
    <body>
        <h1>Translation Report ({source_lang} -> {target_lang})</h1>
        <h2>AI Output</h2>
        <div class="box">{summary}</div>
    </body>
    </html>
    """
    
    h_path = os.path.join(TEMP_DIR, f"Trans_{file_prefix}.html")
    d_path = os.path.join(TEMP_DIR, f"Trans_{file_prefix}.docx")
    
    with open(h_path, "w", encoding="utf-8") as f: f.write(html_out)
    
    # Generate DOCX
    save_as_docx(f"Translation Report ({source_lang} -> {target_lang})", summary, d_path)

    return html_out, d_path, [h_path, d_path]

# --- RESET LOGIC ---
def reset_files(file_list):
    if file_list:
        for f in file_list:
            if os.path.exists(f): 
                try: os.remove(f)
                except: pass
    return None, None, None, None # Reset UI elements

# ==========================================
# GUI LAYOUT (TABS)
# ==========================================
with gr.Blocks(title="AI AV Tools Suite", theme=gr.themes.Soft()) as demo:
    
    gr.Markdown("# 🎬 AI AV Tools Suite")
    
    # Shared State for cleanup
    file_state = gr.State(value=[])

    with gr.Tabs():
        
        # --- TAB 1: MEETING ASSISTANT ---
        with gr.TabItem("📝 Meeting Assistant"):
            with gr.Row():
                with gr.Column():
                    a_file = gr.File(label="Media File")
                    a_key = gr.Textbox(label="Keyword", value="Manu")
                    a_prompt = gr.Textbox(label="Instructions", value="Summarize and list Action Items.", lines=3)
                    a_btn = gr.Button("Generate Report", variant="primary")
                    a_pdf = gr.DownloadButton("Download PDF")
                with gr.Column():
                    a_out = gr.HTML(label="Result")
            
            a_btn.click(process_assistant, [a_file, a_key, a_prompt], [a_out, a_pdf, file_state])

        # --- TAB 2: UNIVERSAL TRANSLATOR ---
        with gr.TabItem("🌐 Universal Translator"):
            with gr.Row():
                with gr.Column():
                    t_file = gr.File(label="Media File")
                    
                    # DYNAMIC LANGUAGE SELECTION
                    with gr.Row():
                        t_source = gr.Dropdown(list(LANG_MAP.keys()), value="Spanish", label="Audio Language (Source)")
                        t_target = gr.Dropdown(["English", "Japanese", "Spanish", "French", "German"], value="Japanese", label="Text Language (Output)")
                    
                    t_btn = gr.Button("Start Translation", variant="primary")
                    t_doc = gr.DownloadButton("Download DOCX")
                with gr.Column():
                    t_out = gr.HTML(label="Translation Result")

            t_btn.click(process_translator, [t_file, t_source, t_target], [t_out, t_doc, file_state])

    # Global Reset
    reset_btn = gr.Button("🗑️ Reset & Cleanup All", variant="stop")
    reset_btn.click(reset_files, inputs=[file_state], outputs=[a_file, t_file, a_out, t_out])

if __name__ == "__main__":
    demo.launch(inbrowser=True)